"""
build_report_docx.py
服务: quantimmu-bench
功能: 读 metrics_ds2.csv + figures_R/*.png, 生成中文 Word 报告
      输出: analysis/BENCHMARK_REPORT.docx
依赖: pip install python-docx
运行: python build_report_docx.py  (从 analysis/ 目录运行, 或脚本会自动定位)
"""
import os, sys
from pathlib import Path
import pandas as pd

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# ── 路径 ──────────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).parent.resolve()
METRICS_CSV = SCRIPT_DIR / "metrics_ds2.csv"
FIGURES_R   = SCRIPT_DIR / "figures_R_v3"
OUT_DOCX    = SCRIPT_DIR / "BENCHMARK_REPORT.docx"

# ── 读数据 ────────────────────────────────────────────────────────────────────
metrics = pd.read_csv(METRICS_CSV)

# ── 辅助函数 ──────────────────────────────────────────────────────────────────
def set_cell_font(cell, text, bold=False, size=10, font_name="SimSun",
                  color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    """设置单元格文字格式"""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.font.name = font_name
    run.font.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 中文字体需同时设 rPr 的 eastAsia
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "SimSun" if not bold else "SimHei")
    rPr.insert(0, rFonts)


def add_heading(doc, text, level=1):
    """添加标题 (黑体)"""
    para = doc.add_heading("", level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = "SimHei"
    run.font.bold = True
    run.font.size = Pt({1: 16, 2: 13, 3: 11}.get(level, 12))
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "SimHei")
    rPr.insert(0, rFonts)


def add_para(doc, text, bold=False, size=10.5, indent=False):
    """添加正文段落 (宋体)"""
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.first_line_indent = Pt(21)
    run = para.add_run(text)
    run.font.name = "SimSun"
    run.font.bold = bold
    run.font.size = Pt(size)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "SimSun")
    rPr.insert(0, rFonts)


def add_bullet(doc, text, size=10.5):
    """添加项目符号段落"""
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "SimSun"
    run.font.size = Pt(size)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "SimSun")
    rPr.insert(0, rFonts)


def add_figure(doc, img_path, caption, width_inch=6.0):
    """嵌入图片 + 图注"""
    if not Path(img_path).exists():
        add_para(doc, f"[图片未找到: {img_path}]")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(img_path), width=Inches(width_inch))
    # 图注
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.font.name = "SimSun"
    cap_run.font.size = Pt(10)
    cap_run.font.italic = True
    r = cap_run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "SimSun")
    rPr.insert(0, rFonts)


def shade_cell(cell, hex_color="D9E1F2"):
    """单元格背景色 (表头用)"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


# ══════════════════════════════════════════════════════════════════════════════
# 开始构建文档
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# 页边距收窄 (1 inch 四边)
from docx.shared import Inches as In
section = doc.sections[0]
section.left_margin   = In(1.0)
section.right_margin  = In(1.0)
section.top_margin    = In(1.0)
section.bottom_margin = In(1.0)

# ── 封面/大标题 ───────────────────────────────────────────────────────────────
add_heading(doc, "新抗原免疫原性预测工具 Benchmark 分析报告", level=1)
add_para(doc, "数据来源：袁老师团队 ELISpot 实验数据（DS1 + DS2）", size=10)
add_para(doc, "分析日期：2026-06-23", size=10)
add_para(doc, "分析工具：DeepImmuno、PredIG、IMPROVE（降级版）、NeoTImmuML（重训版）", size=10)
doc.add_paragraph()

# ── 第 1 节: 背景 ─────────────────────────────────────────────────────────────
add_heading(doc, "1. 背景", level=2)
add_para(doc,
    "本报告对四款新抗原免疫原性预测工具在袁老师团队 ELISpot 实验数据集上的判别能力进行系统评估。"
    "金标准为 ELISpot（酶联免疫斑点法）检测的 T 细胞应答强度（SFC/10⁶细胞）。"
    "DS1 包含 82 条肽段（全阳性，无阴性对照，不能算 AUC）；"
    "DS2 包含 101 条肽段（ELISpot 范围 −33.7 ～ 209 SFC，含 11 条阴性），作为主要评估集。",
    indent=True)
doc.add_paragraph()

# ── 第 2 节: 方法 ─────────────────────────────────────────────────────────────
add_heading(doc, "2. 方法概述", level=2)

add_heading(doc, "2.1 工具与覆盖范围", level=3)
tool_cover = [
    "DeepImmuno：仅覆盖 9-10mer 子肽（DS2 共 11,358 行预测）",
    "PredIG：覆盖 8-14mer 全窗口（DS2 共 34,247 行，最全）",
    "IMPROVE（降级版）：覆盖 8-12mer（DS2 共 26,790 行；跳过 netMHCstabpan，Stability 补 0）",
    "NeoTImmuML（重训版）：覆盖 8-13mer（DS2 共 30,739 行；用 TumorAgDB 数据重训，训练集 364:1 不平衡）",
]
for t in tool_cover:
    add_bullet(doc, t)

add_heading(doc, "2.2 聚合策略", level=3)
add_para(doc,
    "每条全长肽段（Peptide_ID）对应多条子肽预测值，通过三种聚合方式归约为肽段级分数：",
    indent=True)
for t in ["max（最大子肽分）", "mean（所有子肽均值）", "top3mean（最高三条子肽均值）"]:
    add_bullet(doc, t)

add_heading(doc, "2.3 阈值定义与指标", level=3)
add_para(doc, "采用三种 ELISpot 二值化阈值，统计各指标：", indent=True)
for t in [
    "ELISpot > 0（n_pos=90, n_neg=11）",
    "ELISpot > 10（n_pos=64, n_neg=37）",
    "ELISpot > median=22.67（n_pos=50, n_neg=51）",
]:
    add_bullet(doc, t)
add_para(doc, "指标：AUC-ROC（判别能力）、AUPRC（阳性精准-召回曲线下面积）、Spearman ρ（连续相关性）。",
         indent=True)
doc.add_paragraph()

# ── 第 3 节: 结果 ─────────────────────────────────────────────────────────────
add_heading(doc, "3. 结果", level=2)

# 3.1 主要指标表 (agg=max)
add_heading(doc, "3.1 主要指标汇总（DS2，Aggregation = max）", level=3)
add_para(doc, "表 1  各工具 AUC-ROC / AUPRC / Spearman ρ（agg=max，三种阈值）", bold=True, size=10)

tbl_data = metrics[metrics["Aggregation"] == "max"].copy()
tbl_data["Tool"] = pd.Categorical(tbl_data["Tool"],
    categories=["PredIG", "NeoTImmuML", "IMPROVE", "DeepImmuno"], ordered=True)
tbl_data = tbl_data.sort_values(["Tool", "Threshold"])

COLS_DISPLAY = {
    "Tool": "工具",
    "Threshold": "阈值",
    "n_pos": "n_pos",
    "n_neg": "n_neg",
    "AUC_ROC": "AUC-ROC",
    "AUPRC": "AUPRC",
    "Spearman_rho": "Spearman ρ",
    "Spearman_pval": "p 值",
}
display_cols = list(COLS_DISPLAY.keys())

table = doc.add_table(rows=1 + len(tbl_data), cols=len(display_cols))
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头
for j, col in enumerate(display_cols):
    cell = table.cell(0, j)
    shade_cell(cell, "2E75B6")
    set_cell_font(cell, COLS_DISPLAY[col], bold=True, size=9.5,
                  font_name="SimHei", color=(255, 255, 255),
                  align=WD_ALIGN_PARAGRAPH.CENTER)

# 数据行
for i, (_, row) in enumerate(tbl_data.iterrows()):
    bg = "EDF2F8" if i % 6 < 3 else "FFFFFF"
    for j, col in enumerate(display_cols):
        cell = table.cell(i + 1, j)
        shade_cell(cell, bg)
        val = row[col]
        if col in ("AUC_ROC", "AUPRC", "Spearman_rho"):
            text = f"{float(val):.4f}"
        elif col == "Spearman_pval":
            text = f"{float(val):.4f}"
        else:
            text = str(val)
        set_cell_font(cell, text, size=9.5,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# 3.2 聚合方式对比表
add_heading(doc, "3.2 聚合方式对比（DS2，Threshold = ELISpot > 0）", level=3)
add_para(doc, "表 2  各工具三种聚合方式 AUC-ROC 对比", bold=True, size=10)

tbl2_data = metrics[metrics["Threshold"] == ">0"].pivot_table(
    index="Tool", columns="Aggregation", values="AUC_ROC").reset_index()
tbl2_data.columns.name = None
tbl2_data["Tool"] = pd.Categorical(tbl2_data["Tool"],
    categories=["PredIG", "NeoTImmuML", "IMPROVE", "DeepImmuno"], ordered=True)
tbl2_data = tbl2_data.sort_values("Tool")

agg_cols = ["Tool", "max", "mean", "top3mean"]
table2 = doc.add_table(rows=1 + len(tbl2_data), cols=4)
table2.style = "Table Grid"
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_names = ["工具", "max（最大值）", "mean（均值）", "top3mean（Top-3均值）"]
for j, h in enumerate(hdr_names):
    cell = table2.cell(0, j)
    shade_cell(cell, "2E75B6")
    set_cell_font(cell, h, bold=True, size=9.5, font_name="SimHei",
                  color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)

for i, (_, row) in enumerate(tbl2_data.iterrows()):
    for j, col in enumerate(agg_cols):
        cell = table2.cell(i + 1, j)
        shade_cell(cell, "EDF2F8" if i % 2 == 0 else "FFFFFF")
        val = row[col]
        text = f"{float(val):.4f}" if col != "Tool" else str(val)
        set_cell_font(cell, text, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# ── 第 4 节: 图 ───────────────────────────────────────────────────────────────
add_heading(doc, "4. 分析图", level=2)

figures = [
    ("fig1_roc_v3.png",
     "图 1  DS2 ROC 曲线对比（ELISpot > 0，Aggregation = max）。对角虚线为随机基线（AUC=0.50）。"),
    ("fig2_bar_v3.png",
     "图 2  各工具 AUC-ROC 分组柱状图（Aggregation = mean，三种阈值 ELISpot>0/>10/>median 对比）。虚线为随机基线 0.50。"),
    ("fig3_scatter_v3.png",
     "图 3  预测分 vs ELISpot 散点图（Aggregation = max，4 工具分面）。黑线为线性拟合，灰带为 95% CI；角标为 Spearman rho 与 p 值。"),
    ("fig4_bar_v3.png",
     "图 4  三种聚合方式 AUC-ROC 对比（Threshold = ELISpot > 10）。虚线为随机基线 0.50。"),
    ("fig5_heatmap_v3.png",
     "图 5  工具 × 阈值 AUC-ROC 热图（Aggregation = mean）。蓝色越深表示判别能力越强，格内为 AUC 数值。"),
]

for fname, caption in figures:
    img_path = FIGURES_R / fname
    add_figure(doc, img_path, caption, width_inch=6.0)
    doc.add_paragraph()

# ── 第 5 节: 结论 ─────────────────────────────────────────────────────────────
add_heading(doc, "5. 综合结论", level=2)

conclusions = [
    "【PredIG — 第一】最优 AUC-ROC = 0.750（mean 聚合，ELISpot > 0），Spearman ρ = 0.198（p=0.047，显著）；"
    "覆盖最全（8-14mer）；但阈值敏感：严格阈值（>10/>median）时 AUC 下降至 0.56-0.57。",

    "【NeoTImmuML — 第二】AUC-ROC = 0.655（max 聚合），但 Spearman 不显著（ρ=0.022，p=0.829）；"
    "严格阈值（>10）时 AUC 接近随机（0.505）；且为重训版，不代表官方模型。",

    "【IMPROVE — 第三】AUC-ROC = 0.621（max 聚合），为降级版但阈值最稳定；"
    "Spearman 最强（ρ=0.243，p=0.014），说明对 ELISpot 定量强度最敏感；"
    "严格阈值（>10）时 AUC 0.656 反超 PredIG（0.558）。",

    "【DeepImmuno — 第四】所有聚合方式 AUC 均在 0.44-0.52，部分低于随机基线；"
    "Spearman 为负（ρ=−0.117）；仅覆盖 9-10mer，覆盖受限。",
]
for c in conclusions:
    add_bullet(doc, c)

doc.add_paragraph()

# ── 第 6 节: 重要注意事项 (Caveats) ──────────────────────────────────────────
add_heading(doc, "6. 重要注意事项", level=2)

caveats = [
    "样本量极小：DS2 仅 101 条肽段、11 个阴性样本，AUC 置信区间宽，工具排名差异不具统计显著性。",
    "IMPROVE 降级版：跳过 netMHCstabpan 稳定性预测（Stability 补 0），可能压低其实际性能，完整版表现或更优。",
    "NeoTImmuML 重训版：基于 TumorAgDB 重训（训练集正负比 364:1，未做下采样），分数不代表官方模型性能。",
    "DeepImmuno 覆盖受限：仅支持 9-10mer，较短/较长肽段无预测；低 AUC 部分反映表位长度限制而非模型本身。",
    "pTuneos 未纳入：无法直接接受全长肽段输入，需重新设计处理流程，本次不参与比较。",
    "DS1 无阴性对照：DS1 全为阳性（ELISpot 16-677 SFC），无法计算 AUC，四工具 Spearman 均不显著（|ρ| ≤ 0.16）。",
    "AUPRC 高基准线：ELISpot > 0 阈值下 baseline AUPRC = 0.891（=90/101），工具 AUPRC 0.89-0.94 提升有限，"
    "不宜作为主要判别指标。",
]
for i, c in enumerate(caveats, 1):
    add_bullet(doc, f"（{i}）{c}")

doc.add_paragraph()

# ── 第 7 节: 建议下一步 ───────────────────────────────────────────────────────
add_heading(doc, "7. 建议下一步", level=2)

nextsteps = [
    "扩充阴性对照：DS2 仅 11 个阴性，建议增至 n_neg ≥ 30 以获得稳健 AUC 估计。",
    "IMPROVE 完整版对比：补全 netMHCstabpan 运行完整 IMPROVE，与降级版对比分数差异。",
    "PredIG mean 聚合验证：mean vs max 差距 +0.089 AUC，通过交叉验证确认是否稳健。",
    "差异抗原性分析（MT-WT）：利用表中 WT 预测分，计算 MT_score − WT_score，检验净增强是否更能预测 ELISpot。",
    "患者分层分析：按 Patient_ID 归一化 ELISpot，排除患者间系统性偏移的混淆效应。",
    "长度特异评估修正：在子肽层面（不聚合到肽）、按 HLA 类型分组评估，区分 MHC-I 与 MHC-II。",
]
for step in nextsteps:
    add_bullet(doc, step)

doc.add_paragraph()

# ── 保存 ──────────────────────────────────────────────────────────────────────
doc.save(str(OUT_DOCX))
print(f"\nDone. Report saved to: {OUT_DOCX}")
