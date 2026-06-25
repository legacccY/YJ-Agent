# -*- coding: utf-8 -*-
"""
build_report_final.py — 服务 quantimmu-bench
最终交付 Word 报告 (保守版)：余嘉前 5 工具 4 类信息 + 8 工具 ELISpot benchmark 保守结论
                        + 统计稳健性 + DS1 证伪 + 诚实 caveat/许可 + QuantImmune 立项
真源 csv：metrics_ds2_8tools.csv / bootstrap_ci_ds2.csv / bootstrap_paired_ds2.csv
        / ds1_magnitude_spearman_bestbinder.csv (数字经 verifier 0 drift)
措辞红线：保守 (无「最优/最强/无可替代」)；NeoTImmuML 标★自训版；pTuneos CI 用 csv 实值。
运行: python build_report_final.py
"""
import sys
from pathlib import Path
import csv

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

DIR = Path(__file__).parent.resolve()
FIG = DIR / "figures"
FIGD = DIR / "figures_deepdive"
OUT = DIR.parent / "QuantImmuBench_最终交付报告_2026-06-24.docx"


def load_csv(name):
    with open(DIR / name, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


M = load_csv("metrics_ds2_8tools.csv")
BCI = load_csv("bootstrap_ci_ds2.csv")
BP = load_csv("bootstrap_paired_ds2.csv")
DS1 = load_csv("ds1_magnitude_spearman_bestbinder.csv")


def m(tool, agg, thr, col):
    for r in M:
        if r["Tool"] == tool and r["Aggregation"] == agg and r["Threshold"] == thr:
            return r[col]
    return "NA"


def bci(tool, col):
    for r in BCI:
        if r["Tool"] == tool:
            return r[col]
    return "NA"


# ── 字体辅助 ──────────────────────────────────────────────
def _ea(run, font="SimSun"):
    rPr = run._r.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:eastAsia"), font)
    rPr.insert(0, rf)


def heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "SimHei"
    run.font.bold = True
    run.font.size = Pt({1: 17, 2: 14, 3: 12}.get(level, 11))
    run.font.color.rgb = RGBColor(0x0B, 0x3C, 0x49)
    _ea(run, "SimHei")


def para(doc, text, bold=False, size=10.5, indent=False, color=None, italic=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    run.font.name = "SimSun"
    run.font.bold = bold
    run.font.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    _ea(run, "SimSun")
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "SimSun"
    run.font.size = Pt(size)
    _ea(run, "SimSun")


def cell_set(cell, text, bold=False, size=9.5, font="SimSun", color=None,
             align=WD_ALIGN_PARAGRAPH.CENTER, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name = font
    run.font.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    _ea(run, "SimHei" if bold else "SimSun")
    if fill:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)


def make_table(doc, headers, data_rows, colorize_last=False):
    t = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell_set(t.cell(0, j), h, bold=True, size=9.5, color=(255, 255, 255), fill="0B3C49")
    for i, row in enumerate(data_rows):
        bg = "EDF4F4" if i % 2 == 0 else "FFFFFF"
        for j, v in enumerate(row):
            al = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            cell_set(t.cell(i + 1, j), str(v), size=9.5, align=al, fill=bg,
                     bold=(j == 0))
    return t


def figure(doc, path, caption, width=6.2):
    if not Path(path).exists():
        para(doc, f"[图片未找到: {path}]", color=(0xB2, 0x3A, 0x48))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.name = "SimSun"
    r.font.size = Pt(9)
    r.font.italic = True
    _ea(r, "SimSun")


# ══════════════════════════════════════════════════════════
doc = Document()
sec = doc.sections[0]
for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(sec, attr, Inches(0.9))

# 封面
heading(doc, "新抗原免疫原性预测工具 — 部署测试与基准评估报告", level=1)
para(doc, "汇报人：余嘉 (legacccy)  ·  协作项目：袁老师课题组 (癌症个性化新抗原疫苗) · 西交利物浦大学", size=10)
para(doc, "日期：2026-06-24  ·  数据真源：袁老师团队 ELISpot 实验 (DS1 + DS2)；数字经 csv 三方核对 (0 drift)", size=10)
para(doc, "负责工具 (前 5)：PredIG · DeepImmuno · pTuneos · IMPROVE · NeoTImmuML★", size=10, bold=True)

# 摘要
heading(doc, "摘要", level=2)
para(doc,
     "本报告涵盖两部分。(一) 部署测试：在 WSL2 + XJTLU HPC 上部署 5 个新抗原免疫原性预测工具，"
     "逐工具收集 4 类信息 (输入格式 / 运行参数 / 输出含义 / 工具简介)，端到端完整度诚实分三档。"
     "(二) 基准评估：用袁老师 ELISpot 真实数据 (DS2) 对 8 个工具 (前 5 + 第二批 PRIME/ImmuneApp/deepHLApan) "
     "做 benchmark。核心结论保守：现有工具在 ELISpot 上判别力普遍弱、彼此统计不可区分、第二批新工具无增量、"
     "对连续强度 (magnitude) 的定量能力弱。该负结果反而强化自研 QuantImmune 做连续 magnitude 回归的立项动机，"
     "但立项前须先核 IEDB/CEDAR 的连续标签填充率 (命门)。",
     indent=True)

# 1 背景
heading(doc, "1. 任务背景与子任务", level=2)
para(doc,
     "项目总目标：做一个能预测 T 细胞反应「强弱定量程度」的工具 —— 比现有只判「有/无免疫原性」的二分类更进一步。"
     "技术路线 = 大量 benchmark 现有工具 + 数据集，再结合自研 QuantImmune 算法。",
     indent=True)
para(doc, "我 (余嘉) 负责的子任务：在 HPC 部署并测试运行 5 个工具，每个收集以下 4 类信息，最终以 PPT/报告记录：", indent=True)
for t in ["① 输入数据的模板 / 格式", "② 预测工具运行参数 (可调参数的类型及功能)",
          "③ 输出数据格式及含义", "④ 工具的简要介绍 (特点、优势)"]:
    bullet(doc, t)
para(doc, "团队分工：预测工具组 = 余嘉 (前 5) + 李紫晨 (后 5)；QuantImmu 组 = 徐伊琳；数据收集组 = 王子源、谢孟翰。", indent=True)

# 2 部署诚实分级
heading(doc, "2. 5 工具部署状态 (端到端完整度诚实分级)", level=2)
para(doc, "表 1  5 工具部署状态总表", bold=True, size=10)
make_table(doc,
           ["工具", "方法", "能否定量", "本地", "HPC", "端到端完整度 (诚实分级)"],
           [
               ["DeepImmuno", "CNN", "连续0-1", "✅", "✅", "完整端到端 (本地+HPC 双验证)"],
               ["PredIG", "XGBoost", "连续0-1", "✅", "✅", "完整端到端 (本地+HPC 双验证)"],
               ["pTuneos", "ML pipeline", "排名分", "✅", "🟡", "子模型 (Pre&RecNeo 跑肽段, 对账官方 r=1.0)"],
               ["IMPROVE", "RandomForest", "连续0-1", "🟡", "🟡", "降级 (特征链缺 RNA-seq/netMHCstabpan)"],
               ["NeoTImmuML★", "集成 ML", "概率", "🟡", "🟡", "自训版 (无官方权重, 不对标原论文)"],
           ])
para(doc,
     "诚实分级结论：5 工具均部署、均产出可进 ELISpot benchmark 的连续/概率分数，但端到端完整度分三档，不可一概而论。"
     "★ NeoTImmuML 为自训版 —— 官方 repo 无预训练权重，benchmark 用我们用公开数据自训的模型，非官方权重，数值不对标原论文精度。"
     "「⚠️ 降级/部分」= 工具自身限制 (结构性数据缺口)，非部署失败。",
     indent=True, size=10)

# 3 逐工具 4 类信息
heading(doc, "3. 逐工具 4 类信息要点", level=2)
tools_info = [
    ("3.1 DeepImmuno (CNN, 完整端到端)",
     ["输入：CSV 无表头两列 peptide,HLA；肽长死限 9/10-mer；HLA 格式 HLA-A*0201；无需基因组库。",
      "参数：--mode single (单条) / --mode multiple --intdir --outdir (批量)；无可调超参。",
      "输出：tab 分隔 peptide/HLA/immunogenicity，immunogenicity = 连续 0-1 (越高越强)。实测 NLVPMVATV(CMV)=0.957、GILGFVFTL(流感)=0.887，已知强表位高分合理。",
      "简介：最轻量，纯肽+HLA、无许可工具依赖、CPU 可跑。局限：肽长死限 9/10-mer。"]),
    ("3.2 PredIG (XGBoost, 完整端到端)",
     ["输入：3 模式 CSV-Uniprot / CSV-Recombinant / FASTA；肽段 8-14 AA。",
      "参数：--modelXG {neoant|noncan|path} + --type {uniprot|recombinant|fasta} + -o 输出。",
      "输出：CSV，PredIG 列 = 连续 0-1 免疫原性分 + NOAH/NetCleave/物化/TCR_contact 等 13 列特征。实测 SLLMWITQV=0.0061。",
      "简介：连续分 + 可解释特征 + 容器化。部署官方 Docker 14.4G → HPC 转 Singularity 4.6G。局限：镜像大。"]),
    ("3.3 pTuneos (ML pipeline, 子模型端到端)",
     ["输入：完整 pipeline 需 VCF+表达+拷贝数+纯度+HLA (全基因组，吃不了纯肽)；Pre&RecNeo 子模型仅需 MT_pep+WT_pep+HLA 三列 (可跑 ELISpot 肽)。",
      "参数：python pTuneos.py {WES|VCF} -i config.yaml；Pre&RecNeo 经自写 wrapper 调 InVivoModelAndScore。",
      "输出：RefinedNeo = 患者级排名分 (需测序)；Pre&RecNeo (model_pro) = 纯肽免疫原性识别分，与其他工具可比。",
      "简介：镜像自带 netMHCpan-4.0/VEP/GATK/BWA 全套。benchmark 用 Pre&RecNeo 跑 ELISpot 32178 肽对 (对账官方 r=1.0 = 只证复刻逻辑对，非整管线 ELISpot 能力背书)。HPC sif 受限。"]),
    ("3.4 IMPROVE (RandomForest, 降级版)",
     ["输入：TSV 必填 突变肽+野生型肽+HLA；肽段 8-12 AA；两步流程 (feature_calc 算特征 → Predict 跑 RF)。",
      "参数：步2 --model {Simple|TME_excluded|TME_included}；每变体 5 个 RF (rf0-rf4) 集成。",
      "输出：TSV 追加 mean_prediction_rf (连续 0-1) = 5fold×50 RF 集成平均。实测 EEFLNSWML=0.5146。",
      "简介：RF 22 特征。⚠️ 缺口：ELISpot 无 RNA-seq → Expression 特征降级；netMHCstabpan 受 HPC glibc 挡。Predict 步本地+HPC 跑通，非「装不上」。"]),
    ("3.5 NeoTImmuML★ (集成 ML, 自训版)",
     ["输入：CSV = Peptide + immunogenicity(标签) + 78 个 R Peptides 物化特征；肽段 8-13 AA；不要 HLA。",
      "参数：不是 CLI，是 Jupyter notebook (21 cell)；改 file_path 指数据。",
      "输出：分类指标 + 雷达图 + predict_proba 连续概率 → 能定量强弱。★ 官方无权重 → benchmark 用自训版，数值不对标原论文。",
      "简介：纯肽特征、不要 HLA / netMHCpan (部署最轻)。局限：研究 notebook 无预训练权重 → 用 TumorAgDB2.0 重训；不含 78 特征计算代码 (须 R Peptides 算)。"]),
]
for title, lines in tools_info:
    heading(doc, title, level=3)
    for ln in lines:
        bullet(doc, ln, size=10)

# 4 benchmark 方法
heading(doc, "4. ELISpot Benchmark 方法", level=2)
para(doc,
     "测试集 DS2：101 条肽段 (34247 行 = 子肽×HLA)。标签按 ELISpot SFC>0 切：90 阳 / 11 阴 "
     "(11 阴 = 1 个 SFC==0 + 10 个 SFC<0 背景扣减，= 真无反应；阴性定义干净 ≤0)。"
     "参评 8 工具 = 第一批 5 + 第二批 3 (PRIME/ImmuneApp/deepHLApan)。"
     "聚合：每肽多子肽×HLA → max / mean / top3mean；阈值 >0 / >10 / >median=22.67。"
     "主对比口径 = max, >0 (apples-to-apples)。指标 = AUC-ROC、AUPRC、Spearman ρ (连续 SFC 相关)。"
     "AUC 不确定性用 2000 次 bootstrap 95% CI 量化。旧 5 工具复现验证 max|AUC diff| ≤ 0.004，口径对齐 PASS。",
     indent=True)

# 5 主结果
heading(doc, "5. 基准结果 (保守结论)", level=2)
para(doc, "表 2  8 工具主对比 (Aggregation=max, Threshold=ELISpot>0)", bold=True, size=10)
order = ["pTuneos", "PredIG", "NeoTImmuML", "IMPROVE", "ImmuneApp", "PRIME", "DeepImmuno", "deepHLApan"]
rows2 = []
for t in order:
    name = t + "★" if t == "NeoTImmuML" else t
    batch = "第二批(新)" if t in ("PRIME", "ImmuneApp", "deepHLApan") else "第一批"
    rows2.append([name, m(t, "max", ">0", "n_pep"), m(t, "max", ">0", "AUC_ROC"),
                  m(t, "max", ">0", "AUPRC"), m(t, "max", ">0", "Spearman_rho"),
                  m(t, "max", ">0", "Spearman_pval"), batch])
make_table(doc, ["工具", "n_pep", "AUC-ROC", "AUPRC", "Spearman ρ", "ρ p值", "批次"], rows2)
para(doc, "（★ NeoTImmuML 为自训版，非官方权重，不对标原论文精度。PRIME n_pep=100 / deepHLApan n_pep=98 因部分肽段 allele 无分。）", size=9, italic=True)

para(doc, "表 3  各工具最优聚合 AUC (每工具在 9 个聚合×阈值组合里挑最高，不同工具取不同口径，彼此不可比、不作排名依据)", bold=True, size=10)
best = [("pTuneos", "mean", ">0"), ("PredIG", "mean", ">0"), ("IMPROVE", "top3mean", ">10"),
        ("ImmuneApp", "mean", ">0"), ("PRIME", "top3mean", ">median"),
        ("deepHLApan", "top3mean", ">10"), ("DeepImmuno", "mean", ">0")]
rows3 = [[t + ("★" if t == "NeoTImmuML" else ""), f"{a}/{th}", m(t, a, th, "AUC_ROC"),
          m(t, a, th, "Spearman_rho"), m(t, a, th, "Spearman_pval")] for t, a, th in best]
make_table(doc, ["工具", "最优聚合/阈值", "AUC-ROC", "Spearman ρ", "ρ p值"], rows3)

figure(doc, FIG / "fig6_8tools_auc_comparison.png",
       "图 1  8 工具 AUC-ROC (DS2, max 聚合, >0) + 95% bootstrap CI 误差棒；橙色 = 第二批新工具，灰虚线 = 随机基线 0.5。")

para(doc, "三条保守结论：", bold=True)
bullet(doc, f"判别力普遍弱：AUC 点估居前 = pTuneos {bci('pTuneos','AUC')}，但 95% bootstrap CI=[{bci('pTuneos','CI_lo')}, {bci('pTuneos','CI_hi')}] 宽达约 ±0.15；除 pTuneos 外所有工具 CI 下界均跌破随机线 0.5。")
bullet(doc, "统计不可区分：配对 bootstrap 显示 pTuneos vs PredIG ΔAUC=0.091 (CI=[−0.145,+0.310])、vs NeoTImmuML ΔAUC=0.098 (CI=[−0.140,+0.327]) 均跨 0 → 头部工具间分不出统计显著高下。pTuneos 仅显著超弱工具 (IMPROVE/PRIME/deepHLApan)，对最近竞品 PredIG/NeoTImmuML 不显著。根因 = 阴性仅 11 个。")
bullet(doc, "新工具无增量 (稳健结论)：第二批 3 工具在全部聚合×阈值组合下无一超过第一批最优点估 (ImmuneApp 0.589 / PRIME 0.528 / deepHLApan 0.419 低于随机)。该结论不依赖排名精度，且 IEDB 泄漏方向只会让分数虚高、对此结论顺风。")
para(doc, "措辞红线：对外一律用「点估居前」，不用「最优/最强/无可替代」—— 这些一致的数字在 n_neg=11 下不支撑「最优」判语。", italic=True, color=(0xB2, 0x3A, 0x48), size=10)

# 6 统计稳健性
heading(doc, "6. 统计稳健性 (为什么不能下「最优」结论)", level=2)
para(doc, "表 4  8 工具 AUC + 95% bootstrap CI", bold=True, size=10)
rowsb = [[r["Tool"] + ("★" if r["Tool"] == "NeoTImmuML" else ""), r["n_pep"], r["AUC"],
          f"[{r['CI_lo']}, {r['CI_hi']}]", r["CI_width"],
          "是" if r["is_new"] == "1" else "—"] for r in BCI]
make_table(doc, ["工具", "n_pep", "AUC", "95% CI", "CI 宽度", "第二批"], rowsb)
para(doc,
     "根因：少数类 (阴性) 主导 AUC 的不确定性 —— 11 个阴性样本使每次 bootstrap 重抽差异巨大，CI 宽达约 ±0.15，"
     "工具间排名差异 (<0.05 AUC) 完全淹没在 CI 宽度里。pTuneos 0.78 是「单聚合×单阈值×11 阴性」三重最优点："
     "同一工具换 >median 阈值 AUC 掉到约 0.46 (低于随机) → 非稳健能力。",
     indent=True)
figure(doc, FIGD / "fig_bootstrap_ci.png",
       "图 2  8 工具 AUC 点估 + 95% bootstrap CI (caterpillar)；所有 CI 跨越或贴近 0.5 且彼此大幅重叠。")

# 7 定量能力 + DS1
heading(doc, "7. 定量能力 (Spearman) 与 DS1 全阳证伪", level=2)
para(doc,
     "Spearman ρ 衡量工具分数与 ELISpot 连续 SFU 的相关 (= 定量强弱能力)。max 聚合下仅 IMPROVE (ρ=0.243, p=0.014) 与 "
     "PredIG (ρ=0.198, p=0.047) 显著正相关，且为点估居前；切到最优聚合，IMPROVE top3mean ρ=0.320 (p=0.001) 唯一稳定显著。"
     "第二批新工具 Spearman 全部不显著 (|ρ|<0.17, p>0.09)。",
     indent=True)
figure(doc, FIG / "fig7_8tools_spearman.png",
       "图 3  8 工具 Spearman ρ (分数 vs ELISpot SFU, max 聚合)；* = p<0.05。橙色 = 第二批新工具。")

heading(doc, "7.1 DS1 全阳数据集证伪「现有工具能定量」", level=3)
para(doc,
     "DS1 = 82 肽 9mer，全阳性 (ELISpot SFC 范围 16–677，中位 131，无阴性对照 → 算不了 AUC)，幅度跨约 40 倍，"
     "适合测「阳性内部谁更强」的排序。结果：8/9 工具对 DS1 SFC 的 |ρ|<0.16、p 全不显著 (≈随机)。",
     indent=True)
para(doc, "表 5  DS1 全阳数据集肽级 magnitude 排序能力 (best-binder, n=82)", bold=True, size=10)
rowsd = []
for r in DS1:
    if r["tool"] == "NOAH":
        continue
    rho = float(r["spearman_rho"])
    verdict = "显著反向 (非能力)" if abs(rho) > 0.4 else ("弱反向 n.s." if rho < -0.1 else "≈随机")
    rowsd.append([r["tool"], f"{rho:.3f}", f"{float(r['rho_p']):.3g}", verdict])
make_table(doc, ["工具", "Spearman ρ", "p值", "判定"], rowsd)
para(doc,
     "机制：工具判别力主要落在「阳 vs 阴」门槛上；一旦全阳，门槛信息用尽，「阳性内部强弱」基本预测不出。"
     "干净对照 = 同批工具在 DS2 上头部能正向显著排 SFC，到 DS1 全阳子集就全部塌成随机。"
     "对袁课题的正面意义：这是一个诚实的硬结论 —— 现有工具能粗分有无、不能预测连续 magnitude，正坐实 QuantImmune 做连续回归的空白。"
     "（唯一显著的 deepHLApan ρ=−0.50 是反向，为负贡献非能力。）",
     indent=True)

# 8 caveat + 许可
heading(doc, "8. 诚实边界：已知限制与许可红线", level=2)
for c in [
    "样本量极小：DS2 仅 11 个阴性 → 所有 AUC/ρ 的 95% CI 都很宽，工具间排名差异 (<0.05 AUC) 不显著。这是「无稳健最优」结论的直接来源。",
    "患者层聚集：101 肽来自 9 个病人，前 2 个病人贡献约 45% 的阴性肽 → 有效自由度 < 101 (伪重复)，AUC 可能部分在测「区分患者」而非「区分免疫原性」，需按 Patient_ID 分层复核。",
    "IEDB overlap 待测：第二批工具 (PRIME/ImmuneApp/deepHLApan) 多用 IEDB 数据训练，与本 ELISpot 集可能重叠；当前尚无排重代码，「独立性待查」(泄漏方向让分数虚高，对「新工具无增量」主结论顺风，但污染单工具绝对数字)。",
    "工具完整度分级：DeepImmuno/PredIG 完整端到端；pTuneos 子模型 (r=1.0 只证复刻逻辑对)；IMPROVE 特征链降级；NeoTImmuML★ 自训版 (非官方权重)。",
]:
    bullet(doc, c)
para(doc, "⚠️ 许可红线 (强制)：", bold=True, color=(0xB2, 0x3A, 0x48))
bullet(doc, "netMHCpan / netMHCstabpan = DTU 学术许可。第 7(v)/10 条：未经 DTU 书面同意，不得向第三方发布在其软件上跑的 benchmark 结果 (含跑出的数字)。本项目是 benchmark → 论文/对外报告含 netMHCpan 对比数字前须先取 DTU 书面同意 (投稿阶段处理)。")
bullet(doc, "deepHLApan = GPL-2.0，公开发布前需合规审。本报告为内部交付，对外/投稿前须过上述两道许可门。")

# 9 立项
heading(doc, "9. QuantImmune 立项支撑 (蓝海 / 命门 / 天花板 / headline)", level=2)
for title, body in [
    ("蓝海 (方向不撞车)",
     "对 ELISpot SFC 做连续 magnitude 回归是学界公认空白。2024 综述 (Exploration of Immunology, DOI 10.37349/ei.2023.00091) 原文：“Magnitude prediction remains an unaddressed gap.” 逐一核验自称「定量」的工具 (PRIME 自认 ranking 非 magnitude、ICERFIRE 把量级标签塌成二分、neoIM/T-SCAPE 仍做分类) → 真正做连续 magnitude 回归的工具 = 0 个。"),
    ("命门 (立项前零成本必做项)",
     "想做回归却可能没有足够连续标签作 ground truth。绝大多数公开数据集是二元标签 (PRIME/NEPdb/dbPepNeo2 全 binary)；唯一系统性带 magnitude 字段的公开源 = IEDB 及其癌症子库 CEDAR。动手前必须先核 IEDB/CEDAR 的 magnitude 字段实际填充率 —— 若稀疏且 TESLA 补充表无连续列，则「想做回归但无连续标签」直接塌缩立项。零算力，数日可做。"),
    ("理论天花板 (避免过度承诺)",
     "纯「肽+HLA」输入对 magnitude 的可解释方差被生物学结构性封顶 —— 头号因子 naïve precursor frequency (Jenkins & Moon 2012, PMC3334329) 由宿主 TCR 库决定，无法从肽+HLA 序列推出；叠加 ELISpot 测量噪声 (inter-lab CV 可达 40%) → ρ 天花板粗估 0.4–0.6 (低置信，待真实 benchmark 校准)。IMPROVE ρ≈0.32 已达约 2/3，接近但未触顶 —— 请勿承诺 ρ→0.8 颠覆性增益。"),
    ("headline 押 C3 (临床 top-K 排序增量)",
     "连续模型在 held-out 病人上的 top-K 推荐质量优于二分类 —— 临床只能合成 top-K 肽，排序质量直接等于临床价值。这条不赌破天花板，最现实、最可证伪、最有临床说服力。C1 (坐实纯序列天花板) 当诚实能力刻画，C2 (喂供体 TCR-seq 破天花板) 标探索性 stretch goal，不当主承重。"),
]:
    heading(doc, title, level=3)
    para(doc, body, indent=True)

# 9.5 附录：第 9 工具 HLAthena (presentation proxy)
heading(doc, "附录：第 9 工具 HLAthena (presentation proxy) 补充评测", level=2)
def _load_9():
    with open(DIR / "metrics_ds2_9tools.csv", encoding="utf-8") as f:
        lines = [ln for ln in f if not ln.lstrip().startswith("#")]
    return list(csv.DictReader(lines))
_H = [r for r in _load_9() if r["Tool"] == "HLAthena"]
def _hv(agg, thr, col):
    for r in _H:
        if r["Aggregation"] == agg and r["Threshold"] == thr:
            try: return f"{float(r[col]):.4f}"
            except: return str(r[col])
    return "—"
para(doc, "定位（必读）：", bold=True)
para(doc, "HLAthena 预测 MHC-I 抗原提呈 (presentation)，其论文明确声明不预测免疫原性。故本工具在 ELISpot 免疫原性 benchmark 上仅作 presentation baseline proxy 单列报告，不与 8 个免疫原性工具 apples-to-apples 并列。⚠️ 按袁老师分工 HLAthena 属李紫晨负责，本评测为余嘉超额补全。", indent=True)
para(doc, "结果（DS2，n_pep=100/101，数字核 metrics_ds2_9tools.csv）：", bold=True)
make_table(doc,
    ["聚合", "阈值", "AUC-ROC", "AUPRC", "Spearman ρ", "p 值"],
    [["max", ">0", _hv("max", ">0", "AUC_ROC"), _hv("max", ">0", "AUPRC"), _hv("max", ">0", "Spearman_rho"), _hv("max", ">0", "Spearman_pval")],
     ["mean", ">0", _hv("mean", ">0", "AUC_ROC"), _hv("mean", ">0", "AUPRC"), _hv("mean", ">0", "Spearman_rho"), _hv("mean", ">0", "Spearman_pval")],
     ["top3mean", ">0", _hv("top3mean", ">0", "AUC_ROC"), _hv("top3mean", ">0", "AUPRC"), _hv("top3mean", ">0", "Spearman_rho"), _hv("top3mean", ">0", "Spearman_pval")]])
para(doc, "结论：HLAthena 在 ELISpot 免疫原性上近随机 (AUC ≈ 0.51、全聚合×阈值 AUC 范围 0.49–0.59、Spearman ρ 0.08–0.15 且 p 全 > 0.12 不显著) —— 正面印证「提呈 ≠ 免疫原性」，与其设计定位一致。", indent=True)
para(doc, "覆盖度与工程诚实：", bold=True)
for c in [
    "逐肽覆盖 100/101 (98%)：每肽对全 HLA×window 子肽取 max 聚合，即便部分子肽缺失逐肽分仍稳健。",
    "分块完成 266/336：HPC 登录节点跑 (无专用 CPU qos)，70 个 chunk 多为 length-8 在节点高负载下被 cgroup 内存 kill 未完成；因逐肽覆盖已 98% 且结果近随机，对结论无实质影响。",
    "GCS 死锁绕过：镜像空壳运行时拉模型 401 死锁 → 匿名下 65-allele 模型 + patch fetch_models=false 本地挂载跑通。",
]:
    bullet(doc, c)

# 10 结论
heading(doc, "10. 结论与下一步", level=2)
para(doc, "已达成：", bold=True)
for c in [
    "5 工具全部部署 + 4 类信息全收集 (逐工具文档齐)；2 个完整端到端双验证 (DeepImmuno/PredIG)，其余诚实分级。",
    "8 工具 ELISpot benchmark 完成，保守结论 = 判别力普遍弱、统计不可区分、新工具无增量、定量能力弱。",
    "第 9 工具 HLAthena (presentation proxy) 超额补全：ELISpot 上近随机 (AUC 0.51 / ρ 0.08 n.s.)，印证提呈≠免疫原性，单列不并比 (见附录)。",
    "该负结果反而强化 QuantImmune 做连续 magnitude 回归的立项动机。",
]:
    bullet(doc, c)
para(doc, "下一步：", bold=True)
for c in [
    "立项命门 (优先, 零算力)：核 IEDB/CEDAR 的 magnitude 字段填充率 —— 整个方向的开关。",
    "扩负样本：DS2 仅 11 阴性 → 要袁老师补真实 ELISpot 阴性肽至 n_neg≥30 再重测 bootstrap CI。",
    "IEDB overlap 实测：32178 肽 vs IEDB 全库精确 + 9mer 子串 match，报 overlap%。",
    "袁老师正式数据到位后按各工具输入格式做转换脚本 → 正式测试。",
    "对外许可：投稿/公开前取 DTU 书面同意 (netMHCpan 数字) + deepHLApan GPL 合规审。",
]:
    bullet(doc, c)

doc.save(str(OUT))
print("WROTE", OUT)
